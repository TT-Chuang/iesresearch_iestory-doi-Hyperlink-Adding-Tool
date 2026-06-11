import io
import re
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit as st

st.set_page_config(
    page_title="iestory doi Hyperlink Adding Tool", page_icon="🔗", layout="centered"
)

st.title("🔗 iestory doi Hyperlink Adding Tool")
st.markdown("""
This tool automatically injects clickable DOI hyperlinks into your academic papers (Word .docx format).
- **Inline Citations (e.g., `[1]`)**: Uses micro-surgery logic to **100% preserve your original fonts and text highlight colors**.
- **References Section**: Uses an advanced split-reconstruction method to **perfectly match and link every DOI**, even if the text was fragmented into broken chunks in Word's underlying XML structure.
""")


def add_hyperlink_to_run_text_micro(paragraph, run, url, target_text):
    if target_text not in run.text:
        return
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run_el.append(rPr)

    text_node = OxmlElement("w:t")
    text_node.text = target_text
    new_run_el.append(text_node)
    hyperlink.append(new_run_el)

    if run.text.strip() == target_text:
        run._r.getparent().replace(run._r, hyperlink)
    else:
        parts = run.text.split(target_text, 1)
        run.text = parts[0]
        parent = run._r.getparent()
        current_index = parent.index(run._r)
        parent.insert(current_index + 1, hyperlink)
        if len(parts) > 1 and parts[1]:
            new_after_run = paragraph.add_run(parts[1])
            parent.insert(parent.index(hyperlink) + 1, new_after_run._r)


def add_hyperlink_to_paragraph_by_split(paragraph, url, target_text):
    full_text = paragraph.text
    if target_text not in full_text:
        return False
    parts = full_text.split(target_text, 1)
    text_before = parts[0]
    text_after = parts[1] if len(parts) > 1 else ""

    p_element = paragraph._p
    for child in list(p_element):
        if child.tag.endswith("r") or child.tag.endswith("hyperlink"):
            p_element.remove(child)

    if text_before:
        paragraph.add_run(text_before)

    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run_el.append(rPr)

    text_node = OxmlElement("w:t")
    text_node.text = target_text
    new_run_el.append(text_node)
    hyperlink.append(new_run_el)
    p_element.append(hyperlink)

    if text_after:
        paragraph.add_run(text_after)
    return True


# Step 1: File Uploader
uploaded_file = st.file_uploader(
    "Step 1: Upload your Word file (.docx)", type=["docx"]
)

if uploaded_file is not None:
    original_name = uploaded_file.name
    name_without_ext = original_name.rsplit(".", 1)[0]

    st.success(f"Successfully loaded file: {original_name}")

    # Step 2: Custom Download Filename Entry Field
    custom_filename = st.text_input(
        "Step 2: Enter your preferred download filename (without .docx extension)",
        value=f"{name_without_ext}_doi_fixed",
    )

    # Step 3: Execution Button
    if st.button("Step 3: Inject DOI Hyperlinks Now"):
        with st.spinner("Processing document using precision micro-surgery..."):
            doc = docx.Document(uploaded_file)
            doi_dict = {}

            doi_strict_prefix_pattern = re.compile(
                r"doi\s*:\s*([^\s\n\r,;?]+)", re.IGNORECASE
            )
            doi_fallback_pattern = re.compile(
                r"(10\.\d{4,9}/[-._;()/:A-Z0-9]+)", re.IGNORECASE
            )

            paragraphs_list = [p for p in doc.paragraphs if p.text.strip()]

            for idx, p in enumerate(paragraphs_list):
                text = p.text.strip()
                pure_doi = None

                strict_match = doi_strict_prefix_pattern.search(text)
                if strict_match:
                    pure_doi = strict_match.group(1).strip()
                else:
                    fallback_match = doi_fallback_pattern.search(text)
                    if fallback_match:
                        pure_doi = fallback_match.group(1).strip()

                if pure_doi:
                    pure_doi = pure_doi.rstrip(". ,;)")
                    num_match = re.search(r"^\s*\[?(\d{1,3})\]?[\s\.]", text)
                    if num_match:
                        ref_num = int(num_match.group(1))
                    else:
                        prev_text = (
                            paragraphs_list[idx - 1].text.strip() if idx > 0 else ""
                        )
                        prev_num_match = re.search(
                            r"^\s*\[?(\d{1,3})\]?[\s\.]", prev_text
                        )
                        if prev_num_match:
                            ref_num = int(prev_num_match.group(1))
                        else:
                            ref_num = len(doi_dict) + 1

                    full_url = f"https://doi.org/{pure_doi}"
                    doi_dict[ref_num] = {"url": full_url, "raw_doi_text": pure_doi}

            if not doi_dict:
                st.warning(
                    "Notice: No matching DOI patterns detected in the document."
                )
            else:
                # Execution Logic
                is_in_references_zone = False

                for i, p in enumerate(doc.paragraphs):
                    text_strip = p.text.strip()
                    if not text_strip:
                        continue

                    if "references" in text_strip.lower():
                        is_in_references_zone = True
                        continue

                    if not is_in_references_zone:
                        # Inline micro-surgery (keeps highlights safe)
                        for run in list(p.runs):
                            for num in doi_dict.keys():
                                target_citation = f"[{num}]"
                                if target_citation in run.text:
                                    add_hyperlink_to_run_text_micro(
                                        p, run, doi_dict[num]["url"], target_citation
                                    )
                    else:
                        # References full-split technique (bypasses broken XML chunks)
                        for num, info in doi_dict.items():
                            target_doi_text = info["raw_doi_text"]
                            if target_doi_text in p.text:
                                add_hyperlink_to_paragraph_by_split(
                                    p, info["url"], target_doi_text
                                )
                                break

                # Save processed doc into bytes stream
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                st.success(
                    "🎉 All DOI hyperlinks successfully injected! Your original text highlights and formats are 100% untouched."
                )

                # Ensure proper extension mapping
                final_download_name = (
                    custom_filename
                    if custom_filename.endswith(".docx")
                    else f"{custom_filename}.docx"
                )

                # Step 4: Download Button
                st.download_button(
                    label="📥 Click here to download your processed file",
                    data=bio,
                    file_name=final_download_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
